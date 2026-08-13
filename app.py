import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
import re

# Pengaturan Halaman
st.set_page_config(page_title="Generator RPP Madrasah", layout="wide")

# Konfigurasi Kunci API
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
except Exception as e:
    st.error("Kunci API Gemini belum terdeteksi di Streamlit Secrets.")

# --- FUNGSI UTILITAS WORD ---
def set_cell_background(cell, fill_hex):
    """Mengatur warna latar belakang sel tabel di Word"""
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Mengatur margin/padding dalam sel tabel"""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clean_markdown_inline(text):
    """Menghapus formatting markdown sederhana seperti ** atau *"""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    return text.strip()

def markdown_to_docx(markdown_text):
    """Mengonversi teksMarkdown dari Gemini menjadi file Word (.docx) yang terstruktur dan rapi"""
    doc = Document()
    
    # Pengaturan Margin Dokumen (Normal: 1 inci)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    lines = markdown_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Pengecekan Judul (Heading)
        if line.startswith('# '):
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(line[2:]))
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 51, 102)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            i += 1
        elif line.startswith('## '):
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(line[3:]))
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 102, 153)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            i += 1
        elif line.startswith('### '):
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(line[4:]))
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            
        # Pengecekan Tabel Markdown (| header | header |)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i+1] and '-' in lines[i+1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                # Abaikan baris pemisah markdown "|---|---|"
                if not re.match(r'^\s*\|?\s*:?-+:?\s*(\|?\s*:?-+:?\s*)*\|?\s*$', lines[i]):
                    table_lines.append(lines[i])
                i += 1
                
            if table_lines:
                # Parsing baris tabel
                rows_data = []
                for tline in table_lines:
                    cells = [clean_markdown_inline(c) for c in tline.split('|')]
                    # Hilangkan elemen kosong di awal/akhir jika ada
                    if cells and cells[0] == '':
                        cells.pop(0)
                    if cells and cells[-1] == '':
                        cells.pop()
                    if cells:
                        rows_data.append(cells)
                
                if rows_data:
                    num_rows = len(rows_data)
                    num_cols = max(len(r) for r in rows_data)
                    table = doc.add_table(rows=num_rows, cols=num_cols)
                    table.style = 'Table Grid'
                    
                    for r_idx, row in enumerate(rows_data):
                        for c_idx, cell_value in enumerate(row):
                            if c_idx < num_cols:
                                cell = table.cell(r_idx, c_idx)
                                cell.text = cell_value
                                set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
                                
                                # Styling Header Tabel
                                if r_idx == 0:
                                    set_cell_background(cell, "F2F2F2")
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                            run.font.bold = True
                                            run.font.size = Pt(10.5)
                                else:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Arial'
                                            run.font.size = Pt(10)
                    
                    # Tambahkan spasi setelah tabel
                    sp = doc.add_paragraph()
                    sp.paragraph_format.space_before = Pt(0)
                    sp.paragraph_format.space_after = Pt(6)
        
        # Pengecekan Daftar Poin (List)
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(clean_markdown_inline(line[2:]))
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            i += 1
            
        # Teks Paragraf Biasa
        else:
            p = doc.add_paragraph()
            run = p.add_run(clean_markdown_inline(line))
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            i += 1
            
    # Simpan dokumen ke dalam buffer memori
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- FUNGSI LOGIN ---
def check_password():
    """Mengembalikan nilai True jika pengguna sudah memasukkan password yang benar."""
    def password_entered():
        email_input = st.session_state["email"]
        pass_input = st.session_state["password"]
        
        if email_input in st.secrets["passwords"] and st.secrets["passwords"][email_input] == pass_input:
            st.session_state["password_correct"] = True
            del st.session_state["password"]
        else:
            st.session_state["password_correct"] = False

    if "password_correct" not in st.session_state:
        st.session_state["password_correct"] = False

    if not st.session_state["password_correct"]:
        st.title("🔒 Halaman Login")
        st.write("Silakan masukkan Email dan Password untuk mengakses Generator RPP.")
        
        st.text_input("Email", key="email")
        st.text_input("Password", type="password", key="password")
        st.button("Login", on_click=password_entered, type="primary")
        
        if "password_correct" in st.session_state and st.session_state["password_correct"] == False:
            st.error("Email atau Password salah! Silakan coba lagi.")
        
        return False
    return True

# --- APLIKASI UTAMA ---
if check_password():
    if st.button("🚪 Logout"):
        st.session_state["password_correct"] = False
        st.rerun()

    st.title("🤖 Generator RPP Madrasah AI")
    st.write("Selamat datang! Silakan isi formulir di bawah ini untuk menyusun RPP utuh Anda.")

    # Form Input Data
    with st.expander("📝 BAGIAN A: IDENTITAS (Wajib)", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            nama_madrasah = st.text_input("Nama Madrasah")
            nama_guru = st.text_input("Nama Guru")
            mata_pelajaran = st.text_input("Mata Pelajaran")
            fase_kelas = st.text_input("Fase / Kelas")
        with col2:
            semester = st.text_input("Semester")
            tahun_pelajaran = st.text_input("Tahun Pelajaran")
            alokasi_waktu = st.text_input("Alokasi Waktu")
            jumlah_pertemuan = st.number_input("Jumlah Pertemuan", min_value=1, step=1)

    with st.expander("📚 BAGIAN B, C, & D: KOMPETENSI (Wajib)", expanded=True):
        cp = st.text_area("Capaian Pembelajaran (CP)")
        tp = st.text_area("Tujuan Pembelajaran (TP)")
        materi = st.text_area("Materi Pembelajaran (IKTP/Materi)")

    with st.expander("🛠️ BAGIAN E: MODEL PEMBELAJARAN", expanded=True):
        st.info("💡 TIPS: Kosongkan kolom ini jika Anda ingin AI memberikan 3 rekomendasi model terlebih dahulu (Tahap 1).")
        model_pembelajaran = st.text_input("Model Pembelajaran (Opsional)")

    with st.expander("📌 BAGIAN F: INFORMASI TAMBAHAN (Opsional)", expanded=False):
        karakteristik = st.text_area("Karakteristik Murid")
        media = st.text_area("Media Pembelajaran yang Tersedia")
        lingkungan = st.text_area("Lingkungan Belajar")
        catatan = st.text_area("Catatan Tambahan Guru")

    st.markdown("---")

    data_input = f"""
    INFORMASI AWAL PEMBELAJARAN
    A. IDENTITAS
    Nama Madrasah: {nama_madrasah}
    Nama Guru: {nama_guru}
    Mata Pelajaran: {mata_pelajaran}
    Fase/Kelas: {fase_kelas}
    Semester: {semester}
    Tahun Pelajaran: {tahun_pelajaran}
    Alokasi Waktu: {alokasi_waktu}
    Jumlah Pertemuan: {jumlah_pertemuan}

    B. CAPAIAN PEMBELAJARAN: {cp}
    C. TUJUAN PEMBELAJARAN: {tp}
    D. MATERI PEMBELAJARAN: {materi}
    E. MODEL PEMBELAJARAN: {model_pembelajaran}
    F. INFORMASI TAMBAHAN:
    - Karakteristik Murid: {karakteristik}
    - Media Pembelajaran: {media}
    - Lingkungan Belajar: {lingkungan}
    - Catatan Guru: {catatan}
    """

    if st.button("🚀 Proses RPP Sekarang", type="primary"):
        if not cp or not tp or not materi:
            st.warning("Mohon isi Capaian Pembelajaran, Tujuan Pembelajaran, dan Materi terlebih dahulu.")
        else:
            with st.spinner("AI sedang berpikir dan menyusun dokumen Anda..."):
                model = genai.GenerativeModel('gemini-3-flash-preview')
                
                if model_pembelajaran.strip() == "":
                    prompt = f"""
                    Anda bertindak sebagai Tim Pengembang Kurikulum Madrasah.
                    Berikut adalah data pembelajaran:
                    {data_input}
                    
                    # TAHAP 1 (Model Pembelajaran Kosong)
                    Analisis terlebih dahulu:
                    1. Karakteristik materi.
                    2. Tingkat kompleksitas Tujuan Pembelajaran.
                    3. Karakteristik murid berdasarkan fase dan kelas.
                    
                    Selanjutnya berikan minimal 3 rekomendasi model pembelajaran yang paling sesuai.
                    Untuk setiap model jelaskan:
                    - Alasan pedagogis.
                    - Kelebihan.
                    - Kekurangan.
                    - Alasan model tersebut paling sesuai dengan TP.
                    
                    Setelah itu BERHENTI dan berikan catatan (Pilih Model yang Kamu suka, dan masukkan ke Kolom Model Pembelajaran di atas). Jangan membuat RP.
                    """
                else:
                    prompt = f"""
                    Anda bertindak sebagai Tim Pengembang Kurikulum Madrasah (Ahli Kurikulum, Pedagogi, Guru Berprestasi, Praktisi Kurikulum Berbasis Cinta/KBC).
                    
                    Berikut adalah data pembelajaran:
                    {data_input}
                    
                    # TAHAP 2
                    Lakukan analisis internal terhadap hubungan CP dengan TP, materi, karakteristik murid, model pembelajaran, asesmen, DPL, dan KBC.
                    Lalu susun Rencana Pembelajaran (RP) UTUH menggunakan format TABEL MARKDOWN STANDAR.
                    
                    ## A. IDENTITAS
                    Susun dalam Tabel dan Tulis ulang identitas di atas.
                    
                    ## B. IDENTIFIKASI
                    Susun dalam tabel yang hanya 2 kolom utk masing-masing sub bagian yang memuat: Kesiapan Murid, Materi Pelajaran (Faktual, Konseptual, dll), Dimensi Profil Lulusan (DPL), Topik Panca Cinta (WAJIB memuat Cinta kepada Allah dan Rasul-Nya serta Cinta kepada Ilmu), dan Materi Integrasi KBC secara naratif.
                    
                    ## C. DESAIN PEMBELAJARAN
                    Susun dalam tabel yang hanya 2 kolom utk masing-masing sub bagian yang memuat: Capaian Pembelajaran, Lintas Disiplin Ilmu, Tujuan Pembelajaran, Topik, Praktik Pedagogis, Kemitraan, Lingkungan (fisik, budaya, digital), dan Pemanfaatan Digital beserta fungsi pedagogisnya.
                    
                    ## D. PENGALAMAN BELAJAR
                    Buat sejumlah {jumlah_pertemuan} tabel TERPISAH untuk setiap pertemuan.
                    Format WAJIB:
                    Tabel hanya memiliki dua kolom yaitu: "Tahapan Pembelajaran" dan "Pengalaman Belajar".
                    Tahapan meliputi: Awal (Berkesadaran), Inti: Memahami (Bermakna), Inti: Mengaplikasi (Menggembirakan), Refleksi, Penutup.
                    Pada kolom Pengalaman Belajar tuliskan SATU NARASI UTUH yang memuat aktivitas guru, murid, tujuan pedagogis, hubungan TP, dan integrasi KBC. Jangan pisahkan elemen ini ke dalam kolom berbeda.
                    
                    ## E. ASESMEN
                    Susun dalam tabel yang hanya 2 kolom utk masing-masing sub bagian yang berisi Asesmen Awal, Proses, dan Akhir.
                    
                    ATURAN KHUSUS:
                    - Format TP: KKO + Materi + Penguatan Nilai KBC.
                    - Konsistensi Nilai KBC wajib dijaga dari awal hingga Refleksi.
                    - Hindari kalimat pasif/membosankan seperti "Guru menjelaskan materi". Jelaskan APA, MENGAPA, dan Bagaimana.
                    - Jangan tampilkan validasi internal di hasil akhir.
                    """
                
                try:
                    response = model.generate_content(prompt)
                    st.session_state["hasil_rpp"] = response.text
                except Exception as e:
                    st.error(f"Terjadi kesalahan: {e}")

    # Menampilkan Hasil dan Tombol Unduh
    if "hasil_rpp" in st.session_state:
        st.success("Selesai diproses!")
        st.markdown(st.session_state["hasil_rpp"])
        
        # Buat dokumen Word dari teks hasil AI
        doc_download = markdown_to_docx(st.session_state["hasil_rpp"])
        
        st.markdown("---")
        st.download_button(
            label="📥 Unduh Hasil RPP (.docx)",
            data=doc_download,
            file_name=f"RPP_{mata_pelajaran if mata_pelajaran else 'Madrasah'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
